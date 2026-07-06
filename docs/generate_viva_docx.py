"""
Generate a Word (.docx) file containing viva questions and a project explanation.
Run:
  pip install python-docx
  python docs/generate_viva_docx.py

Creates: docs/CareerSwipe_viva_questions.docx
"""
from docx import Document
from docx.shared import Pt

TEXT_FILE = 'docs/viva_questions_and_explanation.txt'
OUT_DOCX_QUESTIONS = 'docs/CareerSwipe_viva_questions.docx'
OUT_DOCX_EXPLANATION = 'docs/CareerSwipe_viva_explanation.docx'

def read_source():
    with open(TEXT_FILE, 'r', encoding='utf-8') as f:
        return f.read()

def create_doc(text):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    lines = text.split('\n')
    for line in lines:
        if line.strip() == '':
            doc.add_paragraph('')
            continue
        # Treat top-level headings
        if line.endswith('— CareerSwipe Project') or line.startswith('Viva Questions') or line.startswith('Concise Project Explanation'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(14)
            continue
        p = doc.add_paragraph(line)
    doc.save(OUT_DOCX_QUESTIONS)
    print(f'Wrote: {OUT_DOCX_QUESTIONS}')


def create_two_docs(text):
    # Split the source text into questions and explanation parts
    split_marker = 'Concise Project Explanation'
    if split_marker in text:
        parts = text.split(split_marker, 1)
        questions_text = parts[0].strip()
        explanation_text = (split_marker + parts[1]).strip()
    else:
        # Fallback: whole text as questions
        questions_text = text
        explanation_text = ''

    # Questions doc
    doc_q = Document()
    style = doc_q.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    for line in questions_text.split('\n'):
        if line.strip() == '':
            doc_q.add_paragraph('')
            continue
        if line.startswith('Viva Questions') or line.endswith('— CareerSwipe Project'):
            p = doc_q.add_paragraph()
            run = p.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(14)
            continue
        doc_q.add_paragraph(line)
    doc_q.save(OUT_DOCX_QUESTIONS)
    print(f'Wrote: {OUT_DOCX_QUESTIONS}')

    # Explanation doc
    doc_e = Document()
    style = doc_e.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    if explanation_text:
        for line in explanation_text.split('\n'):
            if line.strip() == '':
                doc_e.add_paragraph('')
                continue
            if line.startswith('Concise Project Explanation'):
                p = doc_e.add_paragraph()
                run = p.add_run(line.strip())
                run.bold = True
                run.font.size = Pt(14)
                continue
            doc_e.add_paragraph(line)
    else:
        doc_e.add_paragraph('No explanation content found in source file.')
    doc_e.save(OUT_DOCX_EXPLANATION)
    print(f'Wrote: {OUT_DOCX_EXPLANATION}')

if __name__ == '__main__':
    text = read_source()
    create_doc(text)
