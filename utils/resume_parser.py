import re
import os
import uuid
import pdfplumber
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.colors import HexColor

def normalize_date(date_str):
    """Converts various date formats to YYYY-MM-DD for HTML date inputs."""
    if not date_str:
        return ""
    
    # Clean input
    date_str = date_str.strip().lower()
    
    # Try DD/MM/YYYY or MM/DD/YYYY
    match = re.search(r'(\d{1,2})[-/](\d{1,2})[-/](\d{2,4})', date_str)
    if match:
        d1, d2, year = match.groups()
        if len(year) == 2:
            year = "19" + year if int(year) > 40 else "20" + year
        
        # Heuristic for DD/MM vs MM/DD
        try:
            day, month = (int(d1), int(d2)) if int(d1) > 12 else (int(d1), int(d2))
            return f"{int(year):04d}-{int(month):02d}-{int(day):02d}"
        except:
            return ""

    # Try Month DD, YYYY
    months = {'jan':1, 'feb':2, 'mar':3, 'apr':4, 'may':5, 'jun':6, 'jul':7, 'aug':8, 'sep':9, 'oct':10, 'nov':11, 'dec':12}
    match = re.search(r'([a-z]+)\s+(\d{1,2}),?\s+(\d{4})', date_str)
    if match:
        mon_str, day, year = match.groups()
        mon = months.get(mon_str[:3], 1)
        return f"{int(year):04d}-{int(mon):02d}-{int(day):02d}"

    return ""

# Core skills list for matching (Expanded)
SKILLS_LIST = [
    "python", "java", "sql", "html", "css", "javascript", "react", "flask", "django", "node", "aws",
    "c++", "c#", "php", "ruby", "swift", "kotlin", "go", "rust", "typescript", "angular", "vue",
    "mongodb", "postgresql", "mysql", "docker", "kubernetes", "git", "linux", "machine learning",
    "data science", "nlp", "cloud computing", "azure", "gcp", "tableau", "power bi", "excel",
    "communication", "leadership", "project management", "agile", "scrum", "devops", "testing"
]

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file using pdfplumber."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""

def extract_text_from_docx(docx_path):
    """Extracts text from a .docx file using python-docx."""
    try:
        doc = Document(docx_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Word extraction error: {e}")
        return ""

def convert_docx_to_pdf(docx_path, output_folder):
    """Converts .docx to .pdf using reportlab and python-docx.
    
    This is a pure Python solution that doesn't require LibreOffice.
    """
    try:
        # Extract text from DOCX
        doc = Document(docx_path)
        
        # Create PDF
        base_name = os.path.basename(docx_path).rsplit('.', 1)[0]
        pdf_path = os.path.join(output_folder, f"{base_name}.pdf")
        
        # Build PDF with extracted content
        story = []
        styles = getSampleStyleSheet()
        
        # Custom heading style
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=HexColor('#1f2937'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=10,
            textColor=HexColor('#374151'),
            spaceAfter=6,
            leading=14
        )
        
        current_heading = None
        for para in doc.paragraphs:
            if not para.text.strip():
                story.append(Spacer(1, 0.1 * inch))
                continue
            
            text = para.text.strip()
            # Detect headings by checking if they're short and in uppercase/title case
            if len(text) < 50 and (text.isupper() or text.istitle()) and any(c.isupper() for c in text):
                if story:
                    story.append(Spacer(1, 0.15 * inch))
                story.append(Paragraph(text, heading_style))
                current_heading = text
            else:
                story.append(Paragraph(text, body_style))
        
        # Add tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    story.append(Paragraph(row_text, body_style))
        
        # Generate PDF
        doc_pdf = SimpleDocTemplate(pdf_path, pagesize=letter)
        doc_pdf.build(story)
        
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception as e:
        print(f"DOCX to PDF conversion error: {e}")
        return None

def convert_to_pdf(input_path, output_folder):
    """Converts .doc or .docx to .pdf using pure Python."""
    try:
        base_name = os.path.basename(input_path).rsplit('.', 1)[0]
        ext = input_path.rsplit('.', 1)[-1].lower()
        
        if ext == 'docx':
            pdf_path = convert_docx_to_pdf(input_path, output_folder)
            return pdf_path
        elif ext == 'doc':
            # For .doc files, try to convert using reportlab
            # or return None to fall back to extraction only
            print(f"Note: .doc files have limited support. Converting to text-based PDF.")
            try:
                text = extract_text_from_docx(input_path)  # May work with older .doc format
                if text:
                    pdf_path = os.path.join(output_folder, f"{base_name}.pdf")
                    # Create a simple text PDF
                    story = []
                    styles = getSampleStyleSheet()
                    for line in text.split('\n')[:200]:  # Limit lines
                        if line.strip():
                            story.append(Paragraph(line.strip(), styles['BodyText']))
                    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
                    doc.build(story)
                    return pdf_path if os.path.exists(pdf_path) else None
            except:
                return None
        else:
            return None
    except Exception as e:
        print(f"Conversion error: {e}")
        return None

def process_resume(filepath, upload_folder):
    """Main function to handle resume processing."""
    if not os.path.exists(filepath):
        return None

    ext = filepath.rsplit('.', 1)[-1].lower()
    text = ""
    pdf_path = None

    # Step 1: Extract Text
    if ext == 'docx':
        text = extract_text_from_docx(filepath)
        pdf_path = convert_to_pdf(filepath, upload_folder)
    elif ext == 'doc':
        pdf_path = convert_to_pdf(filepath, upload_folder)
        if pdf_path:
            text = extract_text_from_pdf(pdf_path)
    elif ext == 'pdf':
        text = extract_text_from_pdf(filepath)
        pdf_path = filepath

    print("RAW TEXT:", text[:500] + "..." if text else "EMPTY")
    print("TEXT LENGTH:", len(text))

    # Handle Empty Extraction (Scanned PDF)
    if not text.strip():
        print("EXTRACTION FAILED: Scanned or empty file.")
        return {
            "first_name": "Unknown",
            "last_name": "Candidate",
            "email": "",
            "skills": "",
            "resume_path": pdf_path or filepath
        }

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    text_lower = text.lower()
    
    # Step 2: Extract Name (Robust)
    ignore_keywords = ["curriculum", "vitae", "resume", "page", "profile", "contact", "summary", "experience", "education", "skills", "objective", "header"]
    full_name = "Unknown Candidate"
    
    # Try finding label first
    name_label_match = re.search(r'(?:name|full\s?name|candidate)[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', text)
    if name_label_match:
        full_name = name_label_match.group(1).strip()
    else:
        for line in lines[:25]:
            clean_line = line.strip()
            # Skip common header labels
            if any(kw in clean_line.lower() for kw in ignore_keywords): continue
            
            # Name heuristic: 2-4 capitalized words, no common symbols except dot
            if re.match(r'^[A-Z][a-zA-Z.]+(?:\s+[A-Z][a-zA-Z.]+){1,3}$', clean_line):
                full_name = clean_line
                break
            
            # If name is followed by contact info on same line
            head_match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\s*[|•,-]', clean_line)
            if head_match:
                full_name = head_match.group(1).strip()
                break

    # Step 3: Extract Contact Info (Enhanced)
    email_match = re.search(r'[\w.-]+@[\w.-]+', text)
    email = email_match.group(0) if email_match else ""
    
    # Improved phone regex - support multiple international formats
    phone_patterns = [
        r'\+977[- ]?\d{10}',  # Nepal format
        r'977[- ]?\d{10}',    # Nepal without +
        r'\+\d{1,3}[- ]?\(?\d{1,4}\)?[- ]?\d{1,4}[- ]?\d{1,9}',  # International
        r'\(?(\d{3})\)?[- ]?(\d{3})[- ]?(\d{4})',  # US/Canada format
        r'[0-9]{7,15}'  # Generic digits only (7-15 digits)
    ]
    phone = ""
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            phone = phone_match.group(0).strip()
            break
    
    # Step 3b: Extract Gender (Enhanced - more patterns)
    gender = "Other"
    gender_match = re.search(r'(?:gender|sex)[:\s]+(male|female|other|m|f)', text, re.IGNORECASE)
    if gender_match:
        g = gender_match.group(1).lower()
        gender = "Male" if g in ['male', 'm'] else ("Female" if g in ['female', 'f'] else "Other")
    else:
        # Heuristic detection from pronouns and titles
        female_indicators = ["ms.", "mrs.", "miss", "she/her", "female", "woman", "lady"]
        male_indicators = ["mr.", "he/him", "male", "man", "gentleman"]
        
        female_count = sum(1 for indicator in female_indicators if indicator in text_lower)
        male_count = sum(1 for indicator in male_indicators if indicator in text_lower)
        
        if female_count > male_count:
            gender = "Female"
        elif male_count > female_count:
            gender = "Male"

    # Step 4: Address (Robust)
    address = ""
    # Try label first
    addr_label_match = re.search(r'(?:address|location|lives\s?in)[:\s]+([^\n]{5,100})', text, re.IGNORECASE)
    if addr_label_match:
        address = addr_label_match.group(1).strip()
    else:
        addr_keywords = ['Street', 'St.', 'Lane', 'Ln', 'Road', 'Rd', 'Avenue', 'Ave', 'Drive', 'Dr', 'Kathmandu', 'Nepal', 'Lalitpur', 'Bhaktapur', 'Pokhara', 'City', 'Province', 'State']
        for i, line in enumerate(lines[:40]):
            if any(kw in line for kw in addr_keywords) or re.search(r'\b\d{5,6}\b', line):
                # Check if it's not just a phone number
                if not re.match(r'^\+?[\d\s-]{10,15}$', line.strip()):
                    address = line
                    # Optionally append next line if it looks like part of address
                    if i + 1 < len(lines) and any(kw in lines[i+1] for kw in ['Nepal', 'USA', 'UK', 'City']):
                        address += ", " + lines[i+1]
                    break

    # Step 5: Date of Birth (Improved)
    dob = ""
    dob_label_match = re.search(r'(?:dob|date of birth|birth\s?date|born)[:\s]*([\d\w\s,/-]+)', text, re.IGNORECASE)
    if dob_label_match:
        potential_dob = dob_label_match.group(1).strip().split('\n')[0]
        dob = normalize_date(potential_dob)
    
    if not dob:
        # Fallback search for date that looks like a birth year
        dob_match = re.search(r'(\d{1,2}[-/]\d{1,2}[-/](?:19|20)\d{2})', text)
        if dob_match:
            dob = normalize_date(dob_match.group(0))

    # Step 6: Sections (Education & Experience)
    sections_map = {
        "education": [r"education", r"academic", r"qualifications", r"schooling", r"university", r"college"],
        "experience": [r"experience", r"work history", r"employment", r"professional background", r"projects", r"work experience"],
        "skills": [r"skills", r"technical skills", r"competencies", r"expertise", r"technologies"]
    }
    
    extracted_data = {"education": [], "experience": [], "skills": []}
    current_section = None
    
    for line in lines:
        line_clean = line.strip().lower().strip(':')
        is_header = False
        for sec, patterns in sections_map.items():
            if any(re.fullmatch(p, line_clean) for p in patterns) or (len(line_clean) < 25 and any(p in line_clean for p in patterns) and (line.isupper() or line.endswith(':'))):
                current_section = sec
                is_header = True
                break
        
        if not is_header and current_section:
            extracted_data[current_section].append(line)

    # Clean up Education (Find highest qualification)
    edu_text = "\n".join(extracted_data["education"])
    highest_qual = ""
    degree_patterns = [
        r"ph\.?d", r"master's", r"bachelor's", r"m\.?s\.?c", r"b\.?s\.?c", r"m\.?b\.?a", 
        r"b\.?e", r"b\.?tech", r"m\.?tech", r"b\.?a", r"m\.?a", r"diploma", r"h\.?s\.?e\.?b"
    ]
    for pattern in degree_patterns:
        match = re.search(pattern, edu_text, re.IGNORECASE)
        if match:
            for line in extracted_data["education"]:
                if re.search(pattern, line, re.IGNORECASE):
                    highest_qual = line
                    break
            if highest_qual: break
    
    if not highest_qual and extracted_data["education"]:
        highest_qual = extracted_data["education"][0]

    experience_summary = "\n".join(extracted_data["experience"][:15])
    
    # Step 7: Skills
    found_skills = [s.capitalize() for s in SKILLS_LIST if re.search(rf'\b{re.escape(s)}\b', text_lower)]
    if extracted_data["skills"]:
        additional = [s.strip().capitalize() for s in re.split(r'[,|\n•]', " ".join(extracted_data["skills"])) if 2 < len(s.strip()) < 30]
        for s in additional[:15]:
            if s not in found_skills: found_skills.append(s)
    
    skills_str = ", ".join(found_skills)

    # Step 8: Heuristics
    experience_type = "Full-time"
    if "intern" in text_lower: experience_type = "Internship"
    elif "freelance" in text_lower: experience_type = "Freelance"
    elif not experience_summary.strip(): experience_type = "No Experience"

    career_field = "Other"
    field_keywords = {
        "IT / Software": ["python", "java", "software", "developer", "engineer", "it", "coding", "programming", "react", "node"],
        "Marketing / Sales": ["marketing", "sales", "brand", "advertising", "seo", "digital marketing"],
        "Finance / Accounting": ["finance", "accounting", "bank", "audit", "tax", "ca", "accas"],
        "Healthcare / Medical": ["nurse", "doctor", "medical", "health", "hospital", "mbbs"],
        "Education / Teaching": ["teacher", "professor", "tutor", "education", "school", "lecturer"],
        "Design / Creative": ["design", "graphic", "creative", "ui", "ux", "artist", "photoshop"]
    }
    for field, kws in field_keywords.items():
        if any(kw in text_lower for kw in kws):
            career_field = field
            break

    name_parts = full_name.split()
    first_name = name_parts[0] if len(name_parts) > 0 else "Unknown"
    last_name = " ".join(name_parts[1:]) if len(name_parts) > 1 else "Candidate"

    return {
        "name": full_name,
        "first_name": first_name,
        "last_name": last_name,
        "email": email,
        "phone": phone,
        "address": address,
        "gender": gender,  # Use the enhanced extracted gender
        "dob": dob,
        "education": highest_qual,
        "experience": experience_summary,
        "experience_type": experience_type,
        "career_field": career_field,
        "job_location_type": "Remote" if "remote" in text_lower else ("Anywhere" if any(kw in text_lower for kw in ["anywhere", "relocate"]) else "Local"),
        "desired_roles": next((t.capitalize() for t in ["developer", "engineer", "manager", "analyst", "consultant", "designer", "teacher", "nurse", "accountant"] if t in text_lower), ""),
        "employment_type": experience_type if experience_type != "No Experience" else "Full-time",
        "salary": (re.search(r'(?:salary|expectation|expected).*?(\d+[,.]?\d+)', text_lower).group(1).replace(',', '') if re.search(r'(?:salary|expectation|expected).*?(\d+[,.]?\d+)', text_lower) else ""),
        "availability": "1 Month" if any(kw in text_lower for kw in ["1 month", "one month"]) else ("3 Months" if any(kw in text_lower for kw in ["3 month", "three month"]) else "Immediate"),
        "skills": skills_str,
        "resume_path": pdf_path or filepath
    }

