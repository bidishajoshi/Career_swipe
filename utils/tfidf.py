"""
Resume parsing and job recommendation helpers backed by scikit-learn TF-IDF.
"""
import os
import re
import string
from collections import Counter

from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS, TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


STOP_WORDS = set(ENGLISH_STOP_WORDS)

SKILL_TERMS = {
    "accounting", "adobe", "agile", "analytics", "angular", "aws", "azure",
    "banking", "budgeting", "business analysis", "c", "c++", "c#", "cad",
    "caregiving", "case management", "communication", "content marketing",
    "copywriting", "crm", "css", "customer service", "data analysis",
    "data science", "devops", "django", "docker", "education", "excel",
    "figma", "finance", "flask", "git", "go", "google analytics", "gcp",
    "healthcare", "html", "java", "javascript", "kotlin", "kubernetes",
    "leadership", "linux", "logistics", "machine learning", "marketing",
    "mongodb", "mysql", "negotiation", "node", "nursing", "operations",
    "patient care", "php", "postgresql", "power bi", "project management",
    "python", "quality assurance", "react", "recruiting", "research",
    "risk management", "ruby", "rust", "sales", "scrum", "seo", "sql",
    "statistics", "staad pro", "swift", "tableau", "teaching", "testing",
    "typescript", "ui", "ux", "vue", "writing"
}

QUALIFICATION_TERMS = {
    "bachelor", "master", "phd", "doctorate", "diploma", "certificate",
    "certification", "degree", "mba", "bsc", "ba", "msc", "ma", "high school"
}

EXPERIENCE_TERMS = {
    "internship", "junior", "associate", "mid level", "senior", "lead",
    "manager", "director", "years", "experience", "supervisor", "entry level"
}


def parse_resume(filepath: str) -> str:
    """Extract raw text from PDF, DOCX, or DOC resumes."""
    if not filepath or not os.path.exists(filepath):
        return ""

    ext = filepath.rsplit(".", 1)[-1].lower()
    text = ""

    if ext == "pdf" and PDF_SUPPORT:
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            text = ""
    elif ext == "docx" and DOCX_SUPPORT:
        try:
            doc = DocxDocument(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " ".join(cell.text for cell in row.cells)
        except Exception:
            text = ""
    elif ext == "doc":
        try:
            with open(filepath, "r", errors="ignore") as f:
                text = f.read()
        except Exception:
            text = ""

    return text


def clean_text(text: str) -> list[str]:
    """Lowercase, remove punctuation, tokenize, and remove stop words."""
    if not text:
        return []
    punctuation = string.punctuation.replace("+", "").replace("#", "")
    text = text.lower().translate(str.maketrans(punctuation, " " * len(punctuation)))
    tokens = re.findall(r"[a-z0-9][a-z0-9+#.-]*", text)
    return [token for token in tokens if token not in STOP_WORDS and len(token) > 1]


def preprocess_text(text: str) -> str:
    """Return a normalized string suitable for TF-IDF vectorization."""
    return " ".join(clean_text(text))


def _phrase_in_text(phrase: str, text: str) -> bool:
    return re.search(rf"\b{re.escape(phrase)}\b", text, flags=re.IGNORECASE) is not None


def extract_skills(text: str, extra_terms: str = "") -> list[str]:
    """Extract important skills, technologies, qualifications, and keywords."""
    combined = f"{text or ''} {extra_terms or ''}"
    normalized = " ".join(clean_text(combined))
    found = {term for term in SKILL_TERMS if _phrase_in_text(term, normalized)}
    found.update(term for term in QUALIFICATION_TERMS if _phrase_in_text(term, normalized))
    found.update(term for term in EXPERIENCE_TERMS if _phrase_in_text(term, normalized))
    return sorted(found)


def extract_keywords(text: str, top_n: int = 15) -> list[str]:
    """Extract high-signal keywords from text after preprocessing."""
    tokens = clean_text(text)
    if not tokens:
        return []

    counts = Counter(tokens)
    weighted = {}
    for token, count in counts.items():
        weight = 3 if token in SKILL_TERMS else 1
        weighted[token] = count * weight

    return [
        term for term, _ in sorted(weighted.items(), key=lambda item: item[1], reverse=True)[:top_n]
    ]


def build_resume_profile(seeker, resume_text: str) -> dict:
    """Combine extracted resume text with profile fields for personalized matching."""
    profile_text = " ".join([
        resume_text or "",
        seeker.skills or "",
        seeker.education or "",
        seeker.experience or "",
        seeker.career_field or "",
        seeker.desired_roles or "",
        seeker.job_location_type or "",
        seeker.experience_type or "",
    ])
    skills = extract_skills(profile_text, seeker.skills or "")
    keywords = extract_keywords(profile_text, 25)
    return {
        "text": profile_text,
        "preprocessed_text": preprocess_text(profile_text),
        "skills": skills,
        "keywords": keywords,
    }


def job_to_text(job) -> str:
    """Flatten a job record into text used by the recommender."""
    company_name = job.company.company_name if getattr(job, "company", None) else ""
    return " ".join([
        job.title or "",
        company_name,
        job.description or "",
        job.required_skills or "",
        job.location or "",
        job.job_type or "",
        job.job_location_type or "",
        job.experience_level or "",
        job.experience_required or "",
        job.tags or "",
    ])


def _weighted_document(text: str, skills: list[str]) -> str:
    weighted_skills = " ".join(skills * 3)
    return f"{preprocess_text(text)} {weighted_skills}".strip()


def recommend_jobs_for_resume(seeker, resume_text: str, jobs: list, limit: int | None = None) -> list[dict]:
    """Rank jobs by the new 0-100 Match Score algorithm with proper tie-breakers."""
    if not jobs:
        return []

    profile = build_resume_profile(seeker, resume_text)
    resume_doc = _weighted_document(profile["text"], profile["skills"])
    
    # We still compute similarities as a fallback / detail, but prioritize the match score
    if resume_doc and jobs:
        job_docs = []
        job_skills = {}
        for job in jobs:
            skills = extract_skills(job_to_text(job), job.required_skills or "")
            job_skills[job.id] = skills
            job_docs.append(_weighted_document(job_to_text(job), skills))

        try:
            vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2), min_df=1)
            matrix = vectorizer.fit_transform([resume_doc] + job_docs)
            similarities = cosine_similarity(matrix[0:1], matrix[1:]).flatten()
        except Exception:
            similarities = [0.0] * len(jobs)
            job_skills = {job.id: [] for job in jobs}
    else:
        similarities = [0.0] * len(jobs)
        job_skills = {job.id: [] for job in jobs}

    recommendations = []
    resume_skill_set = set(profile["skills"])

    for i, job in enumerate(jobs):
        similarity = similarities[i] if i < len(similarities) else 0.0
        required_skills = set(job_skills.get(job.id, []))
        matched_skills = sorted(required_skills & resume_skill_set)
        missing_skills = sorted(required_skills - resume_skill_set)
        recommended_skills = missing_skills[:6]

        skill_score = (len(matched_skills) / len(required_skills)) if required_skills else 0

        # Unified 0-100 match score
        match_percentage = calculate_job_match_score(seeker, job, resume_text)

        recommendations.append({
            "job": job,
            "match_percentage": match_percentage,
            "similarity_score": round(float(similarity), 4),
            "matched_skills": matched_skills,
            "missing_skills": missing_skills,
            "recommended_skills": recommended_skills,
            "skill_match_percentage": round(skill_score * 100),
            "is_best_match": match_percentage >= 75,
        })

    from datetime import datetime
    class SortKey:
        def __init__(self, match_percentage, created_at, company_name):
            self.match_percentage = match_percentage
            self.created_at = created_at if created_at else datetime.min
            self.company_name = (company_name or "").lower()

        def __lt__(self, other):
            # Sort by match_percentage descending
            if self.match_percentage != other.match_percentage:
                return self.match_percentage > other.match_percentage
            # Sort by created_at descending (newest first)
            if self.created_at != other.created_at:
                return self.created_at > other.created_at
            # Sort by company name ascending (alphabetically)
            return self.company_name < other.company_name

    ranked = sorted(
        recommendations,
        key=lambda item: SortKey(
            item["match_percentage"],
            item["job"].created_at,
            item["job"].company.company_name if getattr(item["job"], "company", None) else ""
        )
    )
    return ranked[:limit] if limit else ranked


def match_resume_to_job(resume_text: str, job_text: str) -> int:
    """Return a 0-100 match score using scikit-learn TF-IDF and cosine similarity."""
    if not resume_text or not job_text:
        return 0
    vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2), min_df=1)
    matrix = vectorizer.fit_transform([preprocess_text(resume_text), preprocess_text(job_text)])
    return round(float(cosine_similarity(matrix[0:1], matrix[1:]).flatten()[0]) * 100)


def match_location_preference(seeker_location: str, seeker_country: str, job_location: str, job_location_type: str) -> float:
    """
    Calculate location match using TF-IDF cosine similarity.
    Returns a 0-1 score indicating how well job location matches seeker preferences.
    
    Args:
        seeker_location: Seeker's preferred location (e.g., "Kathmandu, Nepal")
        seeker_country: Seeker's country
        job_location: Job location (e.g., "Kathmandu")
        job_location_type: Job location type (Remote, Hybrid, Onsite)
    
    Returns:
        Float between 0 and 1 (0 = no match, 1 = perfect match)
    """
    if not job_location or not job_location_type:
        return 0.0
    
    # Remote jobs typically match any location preference
    if job_location_type.lower() == 'remote':
        return 0.95
    
    # Build seeker location preference text
    seeker_text = f"{seeker_location or ''} {seeker_country or ''}".strip().lower()
    job_text = f"{job_location}".strip().lower()
    
    if not seeker_text or not job_text:
        return 0.5 if job_location_type.lower() == 'hybrid' else 0.0
    
    # Exact match gets highest score
    if seeker_text == job_text or seeker_text in job_text or job_text in seeker_text:
        return 0.98
    
    # Use TF-IDF and cosine similarity for partial matches
    try:
        vectorizer = TfidfVectorizer(analyzer='char', ngram_range=(2, 3), min_df=1)
        matrix = vectorizer.fit_transform([seeker_text, job_text])
        similarity = float(cosine_similarity(matrix[0:1], matrix[1:]).flatten()[0])
        
        # Boost score for hybrid positions
        if job_location_type.lower() == 'hybrid':
            similarity = min(1.0, similarity + 0.15)
        
        return min(1.0, similarity)
    except Exception:
        return 0.0


def match_skills_preference(seeker_skills: str, job_required_skills: str) -> float:
    """
    Calculate skill match using TF-IDF cosine similarity.
    Returns a 0-1 score indicating skill alignment.
    
    Args:
        seeker_skills: Comma-separated or space-separated skills from seeker profile
        job_required_skills: Comma-separated or space-separated required skills for job
    
    Returns:
        Float between 0 and 1 (0 = no overlap, 1 = perfect match)
    """
    if not seeker_skills or not job_required_skills:
        return 0.5  # Neutral score if skills not specified
    
    seeker_skills_list = extract_skills(seeker_skills)
    job_skills_list = extract_skills(job_required_skills)
    
    if not seeker_skills_list or not job_skills_list:
        return 0.5
    
    # Calculate overlap percentage
    matched = set(seeker_skills_list) & set(job_skills_list)
    if not matched:
        # Use TF-IDF for partial matches
        try:
            vectorizer = TfidfVectorizer(stop_words="english", min_df=1)
            matrix = vectorizer.fit_transform([
                " ".join(seeker_skills_list),
                " ".join(job_skills_list)
            ])
            similarity = float(cosine_similarity(matrix[0:1], matrix[1:]).flatten()[0])
            return min(1.0, similarity)
        except Exception:
            return 0.3
    
    # Score based on matched skills vs required skills
    match_score = len(matched) / len(job_skills_list)
    return min(1.0, match_score)


def filter_jobs_by_preferences(seeker, jobs: list, min_location_score: float = 0.4, min_skill_score: float = 0.3) -> list[dict]:
    """
    Filter jobs based on seeker's location and skill preferences using TF-IDF cosine similarity.
    
    Args:
        seeker: Seeker model instance with location and skills preferences
        jobs: List of JobListing instances to filter
        min_location_score: Minimum location match score (0-1) to include job
        min_skill_score: Minimum skill match score (0-1) to include job
    
    Returns:
        Sorted list of dicts with job info and match scores
    """
    if not jobs:
        return []
    
    seeker_location = (seeker.address or "").strip()
    seeker_country = (seeker.country or "").strip()
    seeker_skills = (seeker.skills or "").strip()
    
    filtered_results = []
    
    for job in jobs:
        try:
            # Calculate location match
            location_score = match_location_preference(
                seeker_location,
                seeker_country,
                (job.location or "").strip(),
                (job.job_location_type or "Onsite").strip()
            )
            
            # Calculate skill match (default 0.5 if no skills specified)
            if not seeker_skills or not job.required_skills:
                skill_score = 0.5  # Neutral if skills not specified
            else:
                skill_score = match_skills_preference(
                    seeker_skills,
                    (job.required_skills or "").strip()
                )
            
            # Apply location threshold (STRICT) - always filter by location
            if location_score >= min_location_score:
                # Combined relevance score (weighted average)
                relevance_score = (location_score * 0.4) + (skill_score * 0.6)
                
                filtered_results.append({
                    'job': job,
                    'location_score': round(location_score, 3),
                    'skill_score': round(skill_score, 3),
                    'relevance_score': round(relevance_score, 3),
                    'is_relevant': True,
                })
        except Exception as e:
            # Skip jobs that cause errors during filtering
            print(f"Error filtering job {job.id}: {str(e)}")
            continue
    
    # Sort by relevance score (highest first)
    filtered_results.sort(key=lambda x: x['relevance_score'], reverse=True)
    
    return filtered_results


def parse_custom_skills(skills_str: str) -> set[str]:
    """Parse comma, pipe, or newline separated skill strings into a normalized set."""
    if not skills_str:
        return set()
    tokens = re.split(r'[,\n|]+', skills_str)
    return {t.strip().lower() for t in tokens if t.strip()}


def extract_years_of_experience(text: str) -> int:
    """Extract maximum numerical years of experience from profile text."""
    if not text:
        return 0
    matches = re.findall(r'(\d+)\s*(?:year|yr)', text.lower())
    if matches:
        return max(int(m) for m in matches)
    return 0


def get_education_tier(text: str) -> int:
    """Classify text into education tier (1-5)."""
    if not text:
        return 0
    t = text.lower()
    if any(kw in t for kw in ['phd', 'ph.d', 'doctorate', 'doctor of philosophy']):
        return 5
    if any(kw in t for kw in ['master', 'mba', 'msc', 'm.sc', 'ma', 'm.a', 'postgraduate', 'post-graduate']):
        return 4
    if any(kw in t for kw in ['bachelor', 'bsc', 'b.sc', 'ba', 'b.a', 'btech', 'b.tech', 'degree', 'undergraduate']):
        return 3
    if any(kw in t for kw in ['diploma']):
        return 2
    if any(kw in t for kw in ['high school', 'secondary', 'school']):
        return 1
    return 0


def calculate_job_match_score(seeker, job, resume_text: str = None) -> int:
    """
    Calculates a Job Match Score (0–100) for a job listing based on the seeker's profile.
    
    Weights:
      - Skills match: 50%
      - Experience match: 20%
      - Education match: 15%
      - Location match: 10%
      - Certifications & Keywords match: 5%
    """
    try:
        # 1. Skills match (50%)
        seeker_skills = parse_custom_skills(seeker.skills)
        if resume_text:
            resume_skills = set(extract_skills(resume_text))
            seeker_skills = seeker_skills.union(resume_skills)
            
        job_skills = parse_custom_skills(job.required_skills)
        if not job_skills:
            skills_score = 1.0
        else:
            overlap = seeker_skills & job_skills
            skills_score = len(overlap) / len(job_skills)
            
        # 2. Experience match (20%)
        job_req_years = job.min_experience or 0
        if not job_req_years and job.experience_level:
            el = job.experience_level.lower()
            if 'senior' in el:
                job_req_years = 5
            elif 'mid' in el:
                job_req_years = 3
            elif 'entry' in el or 'junior' in el:
                job_req_years = 0
                
        seeker_years = 0
        if seeker.experience_type == 'fresher':
            seeker_years = 0
        else:
            seeker_years = extract_years_of_experience(seeker.experience)
            if seeker_years == 0:
                if seeker.experience:
                    seeker_years = 2
                elif seeker.experience_type == 'experienced':
                    seeker_years = 1
                    
        if job_req_years == 0:
            exp_score = 1.0
        else:
            if seeker_years >= job_req_years:
                exp_score = 1.0
            else:
                exp_score = seeker_years / job_req_years

        # 3. Education match (15%)
        job_desc = (job.description or "").lower()
        job_title = (job.title or "").lower()
        job_edu_tier = get_education_tier(job_title + " " + job_desc)
        
        seeker_edu_text = f"{seeker.education or ''} {seeker.education_history or ''}"
        seeker_edu_tier = get_education_tier(seeker_edu_text)
        if seeker_edu_tier == 0 and seeker_edu_text.strip():
            seeker_edu_tier = 3
            
        if job_edu_tier == 0:
            edu_score = 1.0
        else:
            if seeker_edu_tier >= job_edu_tier:
                edu_score = 1.0
            else:
                edu_score = seeker_edu_tier / job_edu_tier

        # 4. Location match (10%)
        job_loc_type = (job.job_location_type or "Onsite").lower()
        seeker_loc_type = (seeker.job_location_type or "Onsite").lower()
        
        if job_loc_type == "remote":
            loc_type_score = 1.0
            geo_score = 1.0
        else:
            if job_loc_type == seeker_loc_type:
                loc_type_score = 1.0
            elif "hybrid" in job_loc_type or "hybrid" in seeker_loc_type:
                loc_type_score = 0.7
            else:
                loc_type_score = 0.3
                
            seeker_city = (seeker.address or "").lower()
            seeker_country = (seeker.country or "").lower()
            job_city = (job.location or "").lower()
            
            if not job_city:
                geo_score = 1.0
            elif not seeker_city and not seeker_country:
                geo_score = 0.5
            elif (seeker_city and seeker_city in job_city) or (job_city and job_city in seeker_city) or (seeker_country and seeker_country in job_city):
                geo_score = 1.0
            else:
                geo_score = match_location_preference(seeker.address, seeker.country, job.location, job.job_location_type)
                
        loc_score = (loc_type_score * 0.4) + (geo_score * 0.6)

        # 5. Certifications/Keywords match (5%)
        seeker_certs = (seeker.certifications or "").lower()
        job_tags = (job.tags or "").lower()
        
        seeker_cert_words = set(clean_text(seeker_certs))
        job_tag_words = set(clean_text(job_tags))
        
        seeker_profile_words = set(clean_text(f"{seeker.skills or ''} {seeker.experience or ''} {seeker.certifications or ''}"))
        job_words = set(clean_text(f"{job.title} {job.description}"))
        
        if not seeker_cert_words:
            cert_score = 1.0
        else:
            cert_overlap = len(seeker_cert_words & job_words)
            cert_score = 1.0 if cert_overlap > 0 else 0.5
            
        if not job_tag_words:
            tag_score = 1.0
        else:
            tag_overlap = len(seeker_profile_words & job_tag_words)
            tag_score = 1.0 if tag_overlap > 0 else 0.5
            
        cert_keyword_score = (cert_score * 0.5) + (tag_score * 0.5)

        # Final score out of 100
        final_score = (skills_score * 0.50) + (exp_score * 0.20) + (edu_score * 0.15) + (loc_score * 0.10) + (cert_keyword_score * 0.05)
        return max(0, min(100, round(final_score * 100)))
    except Exception as e:
        print(f"Error calculating job match score: {e}")
        return 50 # Default safe fallback score on any error

